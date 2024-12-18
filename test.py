from docx import Document
import win32com.client as win32
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


api_key = ""
client = ""
if api_key is not None:
    from openai import OpenAI

client = OpenAI(api_key=api_key)

def insert_paragraph_after(paragraph, text):
    # Create a new paragraph element (XML)
    new_paragraph = OxmlElement("w:p")

    # Add text to the new paragraph
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    new_paragraph.append(run)

    # Insert the new paragraph after the current one
    paragraph._element.addnext(new_paragraph)


def create_proposal(description):
    goal_str = 'To develop a comprehensive online platform that connects service providers with end users, facilitating seamless communication, booking, and scheduling of services while ensuring security, scalability, and user satisfaction.'
    scope_str = '''User-Friendly Interface
User Registration/Login
Service Provider Dashboard
End User Features
Admin Features
Payment Integration
API Integration
Real-Time Features
'''

    scope_explantion = '''1.	User-Friendly Interface:
a.	Design an intuitive and user-friendly interface for easy navigation. 
2.	User Authentication and authorization
a.	User Registration/Login 
i.	Service providers and end users should be able to register with the platform using email or social media accounts. 
b.	User Profile 
i.	Each user type (admin, service provider, end-user) should have a dedicated profile with relevant information. 
c.	Role Based Access Controls 
i.	Admins should have access to all platform features, service providers should have access to their dashboard, and end users should have access to search and booking functionalities. 
3.	Service Provider Dashboard
a.	Service Listing
i.	Service providers can create, edit, and delete listings for the services they offer, including details such as service description. 
b.	Communication Management 
i.	Service providers should be able to communicate and confirm bookings received from end users. 
c.	Profile Management 
i.	 Service providers can edit their profiles, including contact information and service details. 
4.	End User Feature
a.	Service Search 
i.	End users should be able to search for services based on categories and location 
b.	Booking and Scheduling 
i.	End users can view service provider profiles, book services, and schedule appointments via direct messages 
c.	Reviews and Ratings 
i.	Users can leave reviews and ratings for service providers based on their experiences.
'''
    print("hello")
    prompt = f"your task is to generate the cover-letter for the given job description  {description} you have to give me objectives like a sentence here is example for objective 'To develop a web-based application for service provider marketplace.' dont include 'objective:' in start it should be just the objective line as shown in example "
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}],
    )
    objective = response.choices[0].message.content

    prompt = f"your task is to generate the cover-letter for the given job description  {description} you have to give me Goal like a sentence here is example for Goal 'To develop a comprehensive online platform that connects service providers with end users, facilitating seamless communication, booking, and scheduling of services while ensuring security, scalability, and user satisfaction.' dont include 'goal:' in start it should be just the objective line as shown in example "
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}],
    )
    goal = response.choices[0].message.content
    word_app = win32.Dispatch('Word.Application')
    word_app.Visible = False

    prompt = f"your task is to generate the scope for the given job description  {description} you have to give me scope  like a some points seperated by new line eache line should not have more then 5 words   here is example for scope '{scope_str}' dont include 'scope:' or any other word  in start it should be just the scope points as shown in example and dont include any symbol like '-' ins start"
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}],
    )
    scope = response.choices[0].message.content
    scope = scope.split('\n')
    # print(scope)

    prompt = f"your task is to generate the scope description for the given job description  {description} and here is scope {scope} you have to give me scope  like a some points seperated by new line  here is example for scope explanation '{scope_explantion}' dont include 'scope:' or any other word  in start it should be just the scope points as shown in example and dont include any symbol like '-' in start"
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}],
    )
    scope_details = response.choices[0].message.content

    prompt = f"your task is to generate the Amount for the given job description  {description} here is the example 'The total amount would be $ 15,000' dont include any other word it should be formatted in same way just change the price nothing else"
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}]
    )
    amount = response.choices[0].message.content
    print(scope)
    print(scope_details)
    word_app = win32.Dispatch('Word.Application')
    word_app.Visible = False

    # Open the document
    doc = word_app.Documents.Open('D:\\Projects\\djangoProject2\\document\\proposal.docx')

    # Loop through all shapes (including text boxes)
    for shape in doc.Shapes:
        if shape.TextFrame.HasText:
            text = shape.TextFrame.TextRange.Text
            if 'For Web & Mobile' in text:
                shape.TextFrame.TextRange.Text = text.replace('For Web & Mobile', "Ateeq")

    # Save and close the document
    doc.Save()
    doc.Close()
    word_app.Quit()
    doc = Document("document/proposal.docx")
    position = 0
    # Iterate over paragraphs in the document
    for i, paragraph in enumerate(doc.paragraphs):
        print(paragraph.text)
        if "Client Name" in paragraph.text:
            paragraph.text = paragraph.text.replace("Client Name", "Shahmeer")
        if "To develop a web-based application for service provider marketplace." in paragraph.text:
            paragraph.text = paragraph.text.replace("To develop a web-based application for service provider marketplace.", objective)
        if goal_str in paragraph.text:
            paragraph.text = paragraph.text.replace(goal_str, goal)
        if 'Project Scopes should be placed here' in paragraph.text:
            print("i am here")
            paragraph.text = paragraph.text.replace('Project Scopes should be placed here', scope[0])
        print(amount)
        if 'The total amount would be AED 15,000' in paragraph.text:
            print('****************************************************************************')
            paragraph.text =paragraph.text.replace('The total amount would be AED 15,000',amount)

            for point in scope[1:]:
                # print("i am adding ",point )
                insert_paragraph_after(paragraph, point)
                # paragraph = doc.paragraphs[-1]  #
                i += 1

        if "Replace the scope_details here" in paragraph.text:
            paragraph.text = paragraph.text.replace("Replace the scope_details here", scope_details[0])
            for point in scope_details[1:]:
                # print("i am adding ",point)
                insert_paragraph_after(paragraph, point)
                # paragraph = doc.paragraphs[-1]  #

    prompt = f"your task is to tell the technology stack  for the front-end  for this job description  {description} you have to give me techstacks like a sentence here is example for Tech stacks  'Frontend: react,typescript,redux' Do not include any backend stack like python , node etc "
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}],
    )
    tech_stacks = response.choices[0].message.content

    prompt = f"your task is to tell the technology stack  for the backend-end  for this job description  {description} you have to give me techstacks like a sentence here is example for Tech stacks  'Backend: python,Node,DotNet' Do not include any backend stack like react , vue , typescript etc "
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}],
    )
    back_end = response.choices[0].message.content

    prompt = f"your task is to tell the DataBase  for this job description  {description} you have to give me database like a sentence here is example for Tech stacks  'Database: MongoDB or PostgreSQL for data storage.' Do not include any non related thing"
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": prompt}],
    )
    data_base = response.choices[0].message.content

    for table in doc.tables:
        # Iterate through rows and cells in the table
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
                if "Frontend: React.js/Angular JS for a dynamic and responsive user interface." in cell.text:
                    print('herre')
                    cell.text = cell.text.replace("Frontend: React.js/Angular JS for a dynamic and responsive user interface.", tech_stacks)
                if "Backend: Laravel/ Django for server-side logic" in cell.text:
                    print('herre')
                    cell.text = cell.text.replace("Backend: Laravel/ Django for server-side logic", back_end)

                if "Database: MongoDB or PostgreSQL for data storage." in cell.text:
                    print('herre')
                    cell.text = cell.text.replace("Database: MongoDB or PostgreSQL for data storage.", data_base)


    # Save the changes
    doc.save('document/proposal_output.docx')
    objective += "Relevant Projects:\n"

    # for project in projects:
    #     proposal += f"- {project['name']}: {project['description']}\n"

    return objective


description = '''
I please need a golf analysis software program produced to accurately determine a rotation point location along the length of a golf club before a swing even begins. During the period, golfers commonly move a club in a manner that creates a rotation point along the length of the club. In essence it divides the length of the club into two sections. The point location can vary significantly from golfer to golfer even for the same club and is a crucial measurement that needs to be found.

Videos recorded using standard 2D cameras are all that are fundamentally needed and will be utilized for this application version, so appropriate video tracking will be needed. While perhaps a bit less demanding overall than more advanced 3D tracking, club movements during the period can nevertheless be more difficult to track than when actually swinging.

Club movements can range from being quicker and/or more abrupt than might first thought, to so small and/or subtle that they are barely noticeable unless specifically looking for them. Being able to accurately track such club movements is important. It is currently planned to have tracking markers (such as tapes or dots) placed on clubs used in recordings. That will hopefully aid in more accurate tracking, as well as help to more accurately determine certain club dimensions toward determining the needed solution(s).

Attempts have been partly successful previously. One was done using R programming, which had very good tracking accuracy using a color tracking method, and some nice visual graphics. But among other issues, it was far too slow to really be usable, even when tracking just a very small number of video frames/times (expected to be the case in many instances). I think Python was used in a couple of other attempts. But while tracking speed was much better, one had worthless tracking accuracy, and another lacked multiple elements that prevented determining any solution(s).

Additional documentation more extensively details the good and bad of previous attempts for reference (including past code), along with elements or features needed in order to achieve application success. After accurate tracking is successful, the needed solution(s) can be accurately determined. While not especially complex mathematically, some details are involved largely because some overall club movement also commonly occurs during the period. And such movements in essence need to be filtered out to arrive at the needed solution(s). Steps to arrive at the needed solution(s) are also described in the additional documentation.

Because overall application success has not yet been accomplished, I have still not been able to positively prove certain things yet. In that regard I am still somewhat at a prototyping stage. So I am next in need of a basic version of the application that can at least function properly, yet be as efficient as possible from time and cost standpoints. I would basically be working with and testing on my desktop. If successful and if it matters toward choosing an appropriate programming language(s) at this time, subsequent versions might potentially be desktop, web, and/or phone types.

There is no set budget, but I will spend little more on any further attempts unless and until remaining unproven elements are better proven first. I can then take things from there. Despite being a long-struggling project, various IP rights are still in play for this application. So an NDA will likely be required before being able to release and/or discuss the more detailed documentation.

I will consider project or hourly rate proposals, though I cannot really consider completely open-ended hourly rates. For any hourly rate proposals, a reliable quote of hours needed would be required before I would be able to commit. Please feel free to send any questions or comments my way, as I could have easily missed one or more even very basic items that might be needed for your consideration. Thank you very much.
'''

create_proposal(description=description)
